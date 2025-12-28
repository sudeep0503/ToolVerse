from docx import Document
from openpyxl import Workbook
from openpyxl.cell.cell import MergedCell
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter


def detect_horizontal_merges(row):
    """
    Detect horizontal merged cells in a Word table row.
    Returns list of (start_col, end_col).
    """
    merges = []
    seen = {}

    for idx, cell in enumerate(row.cells):
        seen.setdefault(id(cell), []).append(idx)

    for indices in seen.values():
        if len(indices) > 1:
            merges.append((indices[0], indices[-1]))

    return merges


def detect_vertical_merge_range(table, r, c):
    """
    Best-effort vertical merge detection.
    """
    base_text = table.cell(r, c).text.strip()
    if not base_text:
        return None

    end = r
    for rr in range(r + 1, len(table.rows)):
        if table.cell(rr, c).text.strip() == "":
            end = rr
        else:
            break

    return (r, end) if end > r else None


def copy_alignment(word_cell):
    h = None
    v = None

    if word_cell.paragraphs:
        h = word_cell.paragraphs[0].alignment

    v = word_cell.vertical_alignment

    return Alignment(
        horizontal={0: "left", 1: "center", 2: "right"}.get(h),
        vertical={1: "top", 2: "center", 3: "bottom"}.get(v)
    )


def copy_font(word_cell):
    if not word_cell.paragraphs:
        return None

    for run in word_cell.paragraphs[0].runs:
        if run.text.strip():
            return Font(
                name=run.font.name or "Calibri",
                size=run.font.size.pt if run.font.size else 11,
                bold=run.font.bold,
                italic=run.font.italic
            )
    return None


def docx_to_excel(docx_path, xlsx_path):
    doc = Document(docx_path)
    wb = Workbook()

    default_sheet = wb.active
    sheet_created = False

    for t_idx, table in enumerate(doc.tables, start=1):
        ws = wb.create_sheet(f"Table_{t_idx}")
        sheet_created = True

        written = set()

        for r in range(len(table.rows)):
            row = table.rows[r]

            # Horizontal merges
            for start, end in detect_horizontal_merges(row):
                ws.merge_cells(
                    start_row=r + 1,
                    start_column=start + 1,
                    end_row=r + 1,
                    end_column=end + 1
                )

            for c in range(len(row.cells)):
                if (r, c) in written:
                    continue

                word_cell = table.cell(r, c)

                # excel_cell = ws.cell(row=r + 1, column=c + 1)
                # excel_cell.value = word_cell.text.strip()
                excel_cell = ws.cell(row=r + 1, column=c + 1)

                # 🚨 Skip non-anchor merged cells
                if isinstance(excel_cell, MergedCell):
                    continue
                
                excel_cell.value = word_cell.text.strip()

                align = copy_alignment(word_cell)
                if align:
                    excel_cell.alignment = align

                font = copy_font(word_cell)
                if font:
                    excel_cell.font = font

                written.add((r, c))

                # Vertical merges
                vmerge = detect_vertical_merge_range(table, r, c)
                if vmerge:
                    r_start, r_end = vmerge
                    ws.merge_cells(
                        start_row=r_start + 1,
                        start_column=c + 1,
                        end_row=r_end + 1,
                        end_column=c + 1
                    )
                    for rr in range(r_start + 1, r_end + 1):
                        written.add((rr, c))

        # Auto column width
        for col in ws.columns:
            col_letter = get_column_letter(col[0].column)
            max_len = max((len(str(cell.value)) for cell in col if cell.value), default=0)
            ws.column_dimensions[col_letter].width = min(max_len + 2, 50)

    # Remove default sheet safely
    if sheet_created:
        wb.remove(default_sheet)
    else:
        wb.create_sheet("Empty")

    wb.save(xlsx_path)


# -------------------------
# USAGE
# -------------------------
if __name__ == "__main__":
    docx_to_excel(
        docx_path=r"Files\Report.docx",
        xlsx_path=r"Report_Excel.xlsx"
    )
